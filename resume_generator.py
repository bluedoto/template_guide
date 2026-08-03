import base64
import io
from docx import Document
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import streamlit as st
from weasyprint import HTML

# Set page configuration
st.set_page_config(page_title="Big Tech Resume Builder", layout="wide")

# Custom CSS for styling export buttons and centering headers
st.markdown(
    """
    <style>
    .centered-header {
        text-align: center;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Centered Subtitle / Hook Header
st.markdown(
    "<h1 class='centered-header'>Resume Template that got 1M people into"
    " FAANG</h1>",
    unsafe_allow_html=True,
)
st.markdown(
    "<p class='centered-header'>Good luck! This resume format is solid to get"
    " you anywhere.</p>",
    unsafe_allow_html=True,
)
st.markdown(
    "<p class='centered-header'>Download as either PDF or docx to edit.</p>",
    unsafe_allow_html=True,
)

# --- 1. USER INPUTS (Initialized empty for placeholders) ---
st.header("1. Personal Information")
col1, col2, col3 = st.columns(3)
with col1:
    name = st.text_input("Full Name", "", placeholder="Jane Doe")
with col2:
    email = st.text_input("Email", "", placeholder="jane.doe@example.com")
with col3:
    linkedin = st.text_input(
        "LinkedIn URL", "", placeholder="linkedin.com/in/janedoe"
    )

# --- EDUCATION SECTION (DYNAMIC) ---
st.header("2. Education")
if "education_list" not in st.session_state:
    st.session_state.education_list = [{
        "school": "",
        "location": "",
        "degree": "",
        "major": "",
        "date": "",
    }]

for i, edu in enumerate(st.session_state.education_list):
    display_idx = i + 1
    with st.expander(
            f"Education {display_idx}: {edu.get('school', '') or 'New Entry'}",
            expanded=True,
    ):
        col_edu1, col_edu2 = st.columns(2)
        with col_edu1:
            edu["school"] = st.text_input(
                f"University / Institution {display_idx}",
                edu.get("school", ""),
                placeholder="e.g. Stanford University",
                key=f"edu_school_{display_idx}",
            )
        with col_edu2:
            edu["location"] = st.text_input(
                f"Location {display_idx}",
                edu.get("location", ""),
                placeholder="e.g. Stanford, CA",
                key=f"edu_location_{display_idx}",
            )

        col_edu3, col_edu4 = st.columns(2)
        with col_edu3:
            edu["degree"] = st.text_input(
                f"Degree {display_idx}",
                edu.get("degree", ""),
                placeholder="e.g. Bachelor of Science",
                key=f"edu_degree_{display_idx}",
            )
        with col_edu4:
            edu["major"] = st.text_input(
                f"Major / Focus {display_idx}",
                edu.get("major", ""),
                placeholder="e.g. Computer Science",
                key=f"edu_major_{display_idx}",
            )

        edu["date"] = st.text_input(
            f"Dates {display_idx}",
            edu.get("date", ""),
            placeholder="e.g. 2018 – 2022",
            key=f"edu_date_{display_idx}",
        )

        if len(st.session_state.education_list) > 1:
            if st.button(
                    f"Remove Education {display_idx}", key=f"remove_edu_{display_idx}"
            ):
                st.session_state.education_list.pop(i)
                st.rerun()

if st.button("➕ Add Another Education"):
    st.session_state.education_list.append({
        "school": "",
        "location": "",
        "degree": "",
        "major": "",
        "date": "",
    })
    st.rerun()


# --- PROFESSIONAL EXPERIENCE SECTION (DYNAMIC) ---
st.header("3. Professional Experience")
if "exp_list" not in st.session_state:
    st.session_state.exp_list = [{
        "company": "",
        "location": "",
        "roles": [{"title": "", "subgroup": "", "date": "", "desc": ""}],
    }]

for i, exp in enumerate(st.session_state.exp_list):
    display_idx = i + 1
    with st.expander(
            f"Experience {display_idx}: {exp.get('company', '') or 'New Entry'}",
            expanded=True,
    ):
        col_exp1, col_exp2 = st.columns(2)
        with col_exp1:
            exp["company"] = st.text_input(
                f"Company Name {display_idx}",
                exp.get("company", ""),
                placeholder="e.g. Apple",
                key=f"exp_company_{display_idx}",
            )
        with col_exp2:
            exp["location"] = st.text_input(
                f"Location {display_idx}",
                exp.get("location", ""),
                placeholder="e.g. Cupertino, CA",
                key=f"exp_location_{display_idx}",
            )

        st.markdown("#### Roles / Promotions at this Company")

        if "roles" not in exp:
            exp["roles"] = [{"title": "", "subgroup": "", "date": "", "desc": ""}]

        for r_idx, role in enumerate(exp["roles"]):
            r_display_idx = r_idx + 1
            st.markdown(f"**Role {r_display_idx}**")
            role["title"] = st.text_input(
                f"Job Title {display_idx}-{r_display_idx}",
                role.get("title", ""),
                placeholder="e.g. Software Engineer",
                key=f"role_title_{display_idx}_{r_idx}",
            )
            role["subgroup"] = st.text_input(
                f"Subgroup / Team {display_idx}-{r_display_idx}",
                role.get("subgroup", ""),
                placeholder="e.g. Core OS",
                key=f"role_subgroup_{display_idx}_{r_idx}",
            )
            role["date"] = st.text_input(
                f"Dates Worked for Role {display_idx}-{r_display_idx}",
                role.get("date", ""),
                placeholder="e.g. 06/2022 – Present",
                key=f"role_date_{display_idx}_{r_idx}",
            )
            role["desc"] = st.text_area(
                f"Description / Bullets {display_idx}-{r_display_idx}",
                role.get("desc", ""),
                placeholder=(
                    "• Developed and optimized low-level system daemons...\n•"
                    " Partnered with cross-functional teams..."
                ),
                key=f"role_desc_{display_idx}_{r_idx}",
            )
            if len(exp["roles"]) > 1:
                if st.button(
                        f"Remove Role {r_display_idx}",
                        key=f"remove_role_{display_idx}_{r_idx}",
                ):
                    exp["roles"].pop(r_idx)
                    st.rerun()

        if st.button(
                f"➕ Add Promotion / Role Change for"
                f" {exp.get('company', 'Company')}",
                key=f"add_role_{display_idx}",
        ):
            exp["roles"].append({
                "title": "",
                "subgroup": "",
                "date": "",
                "desc": "",
            })
            st.rerun()

        if len(st.session_state.exp_list) > 1:
            if st.button(
                    f"Remove Entire Experience Block {display_idx}",
                    key=f"remove_exp_{display_idx}",
            ):
                st.session_state.exp_list.pop(i)
                st.rerun()

if st.button("➕ Add Another Experience Company"):
    st.session_state.exp_list.append({
        "company": "",
        "location": "",
        "roles": [{"title": "", "subgroup": "", "date": "", "desc": ""}],
    })
    st.rerun()


# --- 4. OPTIONAL SECTIONS ---
st.header("4. Optional Sections")
col_opt1, col_opt2 = st.columns(2)
with col_opt1:
    include_startup = st.checkbox("Include Startup / Other Experience", value=True)
with col_opt2:
    include_awards = st.checkbox("Include Awards", value=True)

if include_startup:
    st.subheader("Startup / Other Experience")
    if "startup_list" not in st.session_state:
        st.session_state.startup_list = [{
            "company": "",
            "location": "",
            "title": "",
            "date": "",
            "desc": "",
        }]

    for i, startup in enumerate(st.session_state.startup_list):
        display_idx = i + 1
        with st.expander(
                f"Startup {display_idx}: {startup.get('company', '') or 'New Entry'}",
                expanded=True,
        ):
            col_st1, col_st2 = st.columns(2)
            with col_st1:
                startup["company"] = st.text_input(
                    f"Project / Startup Name {display_idx}",
                    startup.get("company", ""),
                    placeholder="e.g. nVidia",
                    key=f"startup_company_{display_idx}",
                )
            with col_st2:
                startup["location"] = st.text_input(
                    f"Location {display_idx}",
                    startup.get("location", ""),
                    placeholder="e.g. Santa Clara, CA",
                    key=f"startup_location_{display_idx}",
                )

            startup["title"] = st.text_input(
                f"Role {display_idx}",
                startup.get("title", ""),
                placeholder="e.g. ML Engineering Intern",
                key=f"startup_title_{display_idx}",
            )
            startup["date"] = st.text_input(
                f"Dates {display_idx}",
                startup.get("date", ""),
                placeholder="e.g. 2020",
                key=f"startup_date_{display_idx}",
            )
            startup["desc"] = st.text_area(
                f"Startup Description {display_idx}",
                startup.get("desc", ""),
                placeholder=(
                    "• Researched and benchmarked neural network inference..."
                ),
                key=f"startup_desc_{display_idx}",
            )
            if len(st.session_state.startup_list) > 1:
                if st.button(
                        f"Remove Startup {display_idx}",
                        key=f"remove_startup_{display_idx}",
                ):
                    st.session_state.startup_list.pop(i)
                    st.rerun()

    if st.button("➕ Add Another Startup/Other Experience"):
        st.session_state.startup_list.append({
            "company": "",
            "location": "",
            "title": "",
            "date": "",
            "desc": "",
        })
        st.rerun()
else:
    if "startup_list" in st.session_state:
        del st.session_state.startup_list

if include_awards:
    st.subheader("Awards & Honors")
    if "awards_list" not in st.session_state:
        st.session_state.awards_list = [{"institution": "", "award": ""}]

    for i, awd in enumerate(st.session_state.awards_list):
        display_idx = i + 1
        with st.expander(f"Award Group {display_idx}", expanded=True):
            col_awd1, col_awd2 = st.columns(2)
            with col_awd1:
                awd["institution"] = st.text_input(
                    f"Awarding Institution / Company {display_idx}",
                    awd.get("institution", ""),
                    placeholder="e.g. Stanford University",
                    key=f"award_inst_{display_idx}",
                )
            with col_awd2:
                pass

            awd["award"] = st.text_input(
                f"Award Description / Title {display_idx}",
                awd.get("award", ""),
                placeholder="e.g. Terman Engineering Scholastic Award",
                key=f"award_text_{display_idx}",
            )
            if len(st.session_state.awards_list) > 1:
                if st.button(
                        f"Remove Award Block {display_idx}",
                        key=f"remove_award_{display_idx}",
                ):
                    st.session_state.awards_list.pop(i)
                    st.rerun()

    if st.button("➕ Add Another Award Block"):
        st.session_state.awards_list.append({"institution": "", "award": ""})
        st.rerun()
else:
    if "awards_list" in st.session_state:
        del st.session_state.awards_list


# --- FALLBACK TO SAMPLE DATA IF INPUTS ARE EMPTY ---
active_name = name.strip() or "Jane Doe"
active_email = email.strip() or "jane.doe@example.com"
active_linkedin = linkedin.strip() or "linkedin.com/in/janedoe"

edu_data = st.session_state.education_list
if len(edu_data) == 1 and not any(edu_data[0].values()):
    active_edu = [{
        "school": "Stanford University",
        "location": "Stanford, CA",
        "degree": "Bachelor of Science",
        "major": "Computer Science",
        "date": "2018 – 2022",
    }]
else:
    active_edu = edu_data

exp_data = st.session_state.exp_list
if (
        len(exp_data) == 1
        and not exp_data[0].get("company")
        and not exp_data[0]["roles"][0].get("title")
):
    active_exp = [
        {
            "company": "Apple",
            "location": "Cupertino, CA",
            "roles": [{
                "title": "Software Engineer",
                "subgroup": "Core Operating Systems",
                "date": "06/2022 – Present",
                "desc": (
                    "• Developed and optimized low-level system daemons,"
                    " improving overall system responsiveness and decreasing"
                    " memory utilization by 14% across next-generation devices.\n•"
                    " Partnered with cross-functional hardware and product"
                    " security teams to architect robust cryptographic validation"
                    " pipelines for system updates."
                ),
            }],
        },
        {
            "company": "Google",
            "location": "Mountain View, CA",
            "roles": [{
                "title": "Software Engineering Intern",
                "subgroup": "Cloud Infrastructure",
                "date": "01/2021 – 05/2022",
                "desc": (
                    "• Built scalable distributed backend services using Go and"
                    " Kubernetes to streamline internal telemetry logging"
                    " frameworks.\n• Reduced query latency for microservices by"
                    " 20% through efficient index restructuring and cache"
                    " optimization strategies."
                ),
            }],
        },
    ]
else:
    active_exp = exp_data

startup_data = st.session_state.get("startup_list", [])
if include_startup and (
        len(startup_data) == 0
        or (len(startup_data) == 1 and not startup_data[0].get("company"))
):
    active_startup = [{
        "company": "nVidia",
        "location": "Santa Clara, CA",
        "title": "Machine Learning Engineering Intern",
        "date": "06/2020 – 09/2020",
        "desc": (
            "• Researched and benchmarked neural network inference performance"
            " on embedded edge-computing accelerator chips.\n• Automated"
            " continuous integration scripts for model compilation benchmarks."
        ),
    }]
else:
    active_startup = startup_data

awards_data = st.session_state.get("awards_list", [])
if include_awards and (
        len(awards_data) == 0
        or (len(awards_data) == 1 and not awards_data[0].get("institution"))
):
    active_awards = [
        {
            "institution": "Stanford University",
            "award": (
                "2022 Terman Engineering Scholastic Award (Top 5% of Senior"
                " Class)"
            ),
        },
        {
            "institution": "Global Hackathon",
            "award": "2021 1st Place Overall Winner out of 1,200+ Global Teams",
        },
    ]
else:
    active_awards = awards_data


# --- 2. HTML PREVIEW FORMATTING ---
def format_bullets(text):
    if not text:
        return ""
    bullets = "".join([
        f"<li>{line.strip('•- ')}</li>"
        for line in text.split("\n")
        if line.strip()
    ])
    return f"<ul class='bullet-list'>{bullets}</ul>"


def format_location_string(comp_name, location):
    if comp_name and location:
        return (
            f"<span class='bold-name'>{comp_name}</span> &middot; <span"
            " style='font-style: italic !important; color: #555555;'"
            f">{location}</span>"
        )
    elif comp_name:
        return f"<span class='bold-name'>{comp_name}</span>"
    return ""


html_template = f"""
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
    @page {{
        size: Letter;
        margin: 0.3in;
    }}
    body {{
        font-family: "Times New Roman", Times, Georgia, serif !important;
        font-size: 10pt;
        color: #000000;
        line-height: 1.25;
        margin: 0px;
    }}
    .header-container {{
        text-align: center;
        margin-bottom: 20px;
    }}
    .name-div {{
        font-family: "Times New Roman", Times, Georgia, serif !important;
        font-size: 20pt;
        font-weight: bold;
        margin-bottom: 4px;
    }}
    .contact-div {{
        font-family: "Times New Roman", Times, Georgia, serif !important;
        font-size: 10pt;
    }}
    .section-header {{
        font-size: 10pt;
        text-transform: uppercase;
        font-weight: bold;
        border-bottom: 1px solid black;
        margin-top: 15px;
        margin-bottom: 8px;
        padding-bottom: 2px;
    }}
    .dark-blue-header {{
        color: #1a4f85; 
        font-size: 10pt;
        margin-top: 10px;
        margin-bottom: 4px;
    }}
    .bold-name {{ font-weight: bold; }}
    .resume-entry {{
        display: table;
        width: 100%;
    }}
    .date-col {{
        display: table-cell;
        width: 160px;
        color: #555555;
        font-style: italic !important;
        vertical-align: top;
        padding-top: 2px;
    }}
    .content-col {{
        display: table-cell;
        vertical-align: top;
    }}
    .bold-text {{ font-weight: bold; display: block; margin-bottom: 3px; }}
    .normal-text {{ font-weight: normal; }}
    .role-block {{
        margin-bottom: 12px;
    }}
    .bullet-list {{
        margin-top: 4px;
        margin-bottom: 8px;
        padding-left: 18px;
    }}
    li {{ margin-bottom: 4px; }}
</style>
</head>
<body>
    <div class="header-container">
        <div class="name-div">{active_name}</div>
        <div class="contact-div">{active_email} | {active_linkedin}</div>
    </div>
    <div class="section-header">EDUCATION</div>
"""

for edu in active_edu:
    header_str = format_location_string(
        edu.get("school", ""), edu.get("location", "")
    )
    d = edu.get("degree", "")
    m = edu.get("major", "")
    if d and m:
        degree_display = (
            f"<span class='bold-text' style='display:inline;'>{d}</span> - <span"
            f" class='normal-text'>{m}</span>"
        )
    elif d:
        degree_display = f"<span class='bold-text' style='display:inline;'>{d}</span>"
    elif m:
        degree_display = f"<span class='normal-text'>{m}</span>"
    else:
        degree_display = ""

    html_template += f"""
    <div class="dark-blue-header">{header_str}</div>
    <div class="resume-entry" style="margin-bottom: 8px;">
        <div class="date-col" style="font-style: italic !important;">{edu.get('date', '')}</div>
        <div class="content-col">{degree_display}</div>
    </div>
    """

html_template += """
    <div class="section-header">PROFESSIONAL EXPERIENCE</div>
"""

for exp in active_exp:
    header_str = format_location_string(
        exp.get("company", ""), exp.get("location", "")
    )
    html_template += f"""
    <div class="dark-blue-header">{header_str}</div>
    """
    for role in exp.get("roles", []):
        subgroup_str = (
            f"<i style='display: block; margin-bottom: 5px; color: #000;'>["
            f"{role.get('subgroup', '')}]</i>"
            if role.get("subgroup")
            else ""
        )

        html_template += f"""
        <div class="resume-entry">
            <div class="date-col" style="font-style: italic !important;">{role.get('date', '')}</div>
            <div class="content-col role-block">
                <span class="bold-text">{role.get('title', '')}</span>
                {subgroup_str}
                {format_bullets(role.get('desc', ''))}
            </div>
        </div>
        """

if include_startup and active_startup:
    html_template += """
    <div class="section-header">OTHER EXPERIENCE</div>
    """
    for startup in active_startup:
        header_str = format_location_string(
            startup.get("company", ""), startup.get("location", "")
        )
        html_template += f"""
        <div class="dark-blue-header">{header_str}</div>
        <div class="resume-entry" style="margin-bottom: 8px;">
            <div class="date-col" style="font-style: italic !important;">{startup.get('date', '')}</div>
            <div class="content-col">
                <span class="bold-text" style="display:inline;">{startup.get('title', '')}</span>
                {format_bullets(startup.get('desc', ''))}
            </div>
        </div>
        """

if include_awards and active_awards:
    html_template += """
    <div class="section-header">AWARDS & HONORS</div>
    """
    for awd in active_awards:
        inst_str = (
            f"<i style='color: #555555;'>{awd.get('institution', '')}</i>"
            if awd.get("institution")
            else ""
        )
        html_template += f"""
        <div class="resume-entry" style="margin-bottom: 4px;">
            <div class="date-col" style="font-style: italic !important;">{inst_str}</div>
            <div class="content-col">{awd.get('award', '')}</div>
        </div>
        """

html_template += """
</body>
</html>
"""


# --- 3. WORD DOCUMENT GENERATION ---
def generate_docx():
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = docx.shared.Pt(10)
    font.color.rgb = docx.shared.RGBColor(0, 0, 0)

    for section in doc.sections:
        section.top_margin = docx.shared.Inches(0.3)
        section.bottom_margin = docx.shared.Inches(0.3)
        section.left_margin = docx.shared.Inches(0.3)
        section.right_margin = docx.shared.Inches(0.3)

    def add_para_run(p, text, bold=False, italic=False, color=None, size_pt=10):
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r.bold = bold
        r.italic = italic
        r.font.size = docx.shared.Pt(size_pt)
        if color:
            r.font.color.rgb = color
        return r

    def add_section_heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = docx.shared.Pt(14)
        p.paragraph_format.space_after = docx.shared.Pt(6)
        r = p.add_run(title)
        r.font.name = "Times New Roman"
        r.bold = True
        r.font.size = docx.shared.Pt(10)

        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "2")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Name and Contact
    p_name = doc.add_paragraph()
    p_name.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = docx.shared.Pt(2)
    add_para_run(p_name, active_name, bold=True, size_pt=20)

    p_contact = doc.add_paragraph()
    p_contact.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = docx.shared.Pt(10)
    add_para_run(p_contact, f"{active_email} | {active_linkedin}", size_pt=10)

    def add_entry_row(date_text, content_callback, is_italic_date=True):
        tbl = doc.add_table(rows=1, cols=2)
        tbl.autofit = False

        col_widths = [docx.shared.Inches(1.6), docx.shared.Inches(5.9)]
        for i, col in enumerate(tbl.columns):
            col.width = col_widths[i]
            for cell in col.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(int(col_widths[i].inches * 1440)))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)

        tblPr = tbl._tbl.tblPr
        tblBorders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            b = OxmlElement(f"w:{border_name}")
            b.set(qn("w:val"), "none")
            tblBorders.append(b)
        tblPr.append(tblBorders)

        cell_date = tbl.cell(0, 0)
        cell_content = tbl.cell(0, 1)

        for cell in [cell_date, cell_content]:
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement("w:tcMar")
            for m in ["top", "bottom", "left", "right"]:
                node = OxmlElement(f"w:{m}")
                node.set(qn("w:w"), "0")
                node.set(qn("w:type"), "dxa")
                tcMar.append(node)
            tcPr.append(tcMar)

        p_date = cell_date.paragraphs[0]
        p_date.paragraph_format.space_after = docx.shared.Pt(4)
        p_date.paragraph_format.space_before = docx.shared.Pt(0)
        add_para_run(
            p_date,
            date_text,
            italic=is_italic_date,
            color=docx.shared.RGBColor(85, 85, 85),
            size_pt=10,
        )

        p_content = cell_content.paragraphs[0]
        p_content.paragraph_format.space_after = docx.shared.Pt(2)
        p_content.paragraph_format.space_before = docx.shared.Pt(0)
        content_callback(cell_content, p_content)

        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = docx.shared.Pt(0)
        p_space.paragraph_format.space_before = docx.shared.Pt(0)

    # EDUCATION
    add_section_heading("EDUCATION")
    for edu in active_edu:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = docx.shared.Pt(6)
        p_sub.paragraph_format.space_after = docx.shared.Pt(2)
        add_para_run(
            p_sub,
            f"{edu.get('school', '')}",
            bold=True,
            color=docx.shared.RGBColor(26, 79, 133),
        )
        if edu.get("location"):
            add_para_run(
                p_sub,
                f" · {edu.get('location')}",
                italic=True,
                color=docx.shared.RGBColor(85, 85, 85),
            )

        def fill_edu_content(cell, p):
            d = edu.get("degree", "")
            m = edu.get("major", "")
            if d:
                add_para_run(p, d, bold=True)
            if m:
                add_para_run(p, f" - {m}", bold=False)

        add_entry_row(edu.get("date", ""), fill_edu_content)

    # PROFESSIONAL EXPERIENCE
    add_section_heading("PROFESSIONAL EXPERIENCE")
    for exp in active_exp:
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = docx.shared.Pt(8)
        p_sub.paragraph_format.space_after = docx.shared.Pt(2)
        add_para_run(
            p_sub,
            f"{exp.get('company', '')}",
            bold=True,
            color=docx.shared.RGBColor(26, 79, 133),
        )
        if exp.get("location"):
            add_para_run(
                p_sub,
                f" · {exp.get('location')}",
                italic=True,
                color=docx.shared.RGBColor(85, 85, 85),
            )

        for role in exp.get("roles", []):

            def fill_exp_content(cell, p_first):
                add_para_run(p_first, role.get("title", ""), bold=True)
                if role.get("subgroup"):
                    p_subg = cell.add_paragraph()
                    p_subg.paragraph_format.space_after = docx.shared.Pt(4)
                    add_para_run(
                        p_subg, f"[{role.get('subgroup')}]", italic=True
                    )

                for line in role.get("desc", "").split("\n"):
                    if line.strip():
                        p_b = cell.add_paragraph(style="List Bullet")
                        p_b.paragraph_format.space_after = docx.shared.Pt(3)
                        p_b.paragraph_format.space_before = docx.shared.Pt(0)
                        add_para_run(p_b, line.strip("•- "))

            add_entry_row(role.get("date", ""), fill_exp_content)

    # OTHER EXPERIENCE
    if include_startup and active_startup:
        add_section_heading("OTHER EXPERIENCE")
        for startup in active_startup:
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.space_before = docx.shared.Pt(8)
            p_sub.paragraph_format.space_after = docx.shared.Pt(2)
            add_para_run(
                p_sub,
                f"{startup.get('company', '')}",
                bold=True,
                color=docx.shared.RGBColor(26, 79, 133),
            )
            if startup.get("location"):
                add_para_run(
                    p_sub,
                    f" · {startup.get('location')}",
                    italic=True,
                    color=docx.shared.RGBColor(85, 85, 85),
                )

            def fill_startup_content(cell, p_first):
                add_para_run(p_first, startup.get("title", ""), bold=True)
                for line in startup.get("desc", "").split("\n"):
                    if line.strip():
                        p_b = cell.add_paragraph(style="List Bullet")
                        p_b.paragraph_format.space_after = docx.shared.Pt(3)
                        p_b.paragraph_format.space_before = docx.shared.Pt(0)
                        add_para_run(p_b, line.strip("•- "))

            add_entry_row(startup.get("date", ""), fill_startup_content)

    # AWARDS & HONORS
    if include_awards and active_awards:
        add_section_heading("AWARDS & HONORS")

        tbl = doc.add_table(rows=len(active_awards), cols=2)
        tbl.autofit = False

        col_widths = [docx.shared.Inches(1.6), docx.shared.Inches(5.9)]
        for i, col in enumerate(tbl.columns):
            col.width = col_widths[i]
            for cell in col.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                tcW.set(qn("w:w"), str(int(col_widths[i].inches * 1440)))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)

        tblPr = tbl._tbl.tblPr
        tblBorders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            b = OxmlElement(f"w:{border_name}")
            b.set(qn("w:val"), "none")
            tblBorders.append(b)
        tblPr.append(tblBorders)

        for idx, awd in enumerate(active_awards):
            cell_inst = tbl.cell(idx, 0)
            cell_awd = tbl.cell(idx, 1)

            for cell in [cell_inst, cell_awd]:
                tcPr = cell._tc.get_or_add_tcPr()
                tcMar = OxmlElement("w:tcMar")
                for m in ["top", "bottom", "left", "right"]:
                    node = OxmlElement(f"w:{m}")
                    node.set(qn("w:w"), "0")
                    node.set(qn("w:type"), "dxa")
                    tcMar.append(node)
                tcPr.append(tcMar)

            p_inst = cell_inst.paragraphs[0]
            p_inst.paragraph_format.space_after = docx.shared.Pt(4)
            p_inst.paragraph_format.space_before = docx.shared.Pt(0)
            add_para_run(
                p_inst,
                awd.get("institution", ""),
                italic=True,
                color=docx.shared.RGBColor(85, 85, 85),
                size_pt=10,
            )

            p_awd = cell_awd.paragraphs[0]
            p_awd.paragraph_format.space_after = docx.shared.Pt(4)
            p_awd.paragraph_format.space_before = docx.shared.Pt(0)
            add_para_run(p_awd, awd.get("award", ""))

        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_after = docx.shared.Pt(0)
        p_space.paragraph_format.space_before = docx.shared.Pt(0)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


# --- 4. PREVIEW & EXPORT OPTIONS ---
st.header("Preview")

# Corrected to pass the HTML string cleanly to st.iframe
st.iframe(src=html_template, height=800, scrolling=True)

# Generate PDF using Weasyprint
pdf_bytes = HTML(string=html_template).write_pdf()

st.header("Download Options")
col_d1, col_d2 = st.columns(2)

with col_d1:
    st.download_button(
        label="📥 Download PDF Resume",
        data=pdf_bytes,
        file_name=f"{active_name.replace(' ', '_')}_Resume.pdf",
        mime="application/pdf",
        use_container_width=True,
    )

with col_d2:
    docx_bytes = generate_docx()
    st.download_button(
        label="📥 Download Word (.docx) Resume",
        data=docx_bytes,
        file_name=f"{active_name.replace(' ', '_')}_Resume.docx",
        mime=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        use_container_width=True,
    )